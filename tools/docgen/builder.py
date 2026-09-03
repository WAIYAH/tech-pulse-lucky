"""
TechPulseDocument - the house style for every Get Techy With Lucky learning document.

Wraps python-docx so content modules describe *what* a document says while this
module owns *how* it looks: cover page, running header, footer with "Page X of Y",
a light watermark, the type scale, and the standard instructional blocks
(chapters, terminology tables, code, callouts, quizzes, assignments).
"""

from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from . import brand as B


# ==========================================================================
# Low-level OOXML helpers
# ==========================================================================

def set_run_font(run, name: str) -> None:
    """Apply a font to every script slot so Word never substitutes silently."""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def shade_paragraph(paragraph, hex_fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    ppr.append(shd)


def shade_cell(cell, hex_fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcpr.append(shd)


def paragraph_border(paragraph, *, edges=("bottom",), hex_color=B.HEX_RULE, size=6, space=6) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for edge in edges:
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), str(space))
        el.set(qn("w:color"), hex_color)
        borders.append(el)
    ppr.append(borders)


def clear_inherited_tabs(paragraph, positions=(4680, 9360)) -> None:
    """
    Paragraph tab stops MERGE with the ones inherited from the style rather than
    replacing them, so Word's Letter-width Footer stops would still win. OOXML's
    only way to drop an inherited stop is an explicit w:val="clear" at its position.
    """
    ppr = paragraph._p.get_or_add_pPr()
    tabs = ppr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        ppr.append(tabs)
    for index, pos in enumerate(positions):
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "clear")
        tab.set(qn("w:pos"), str(pos))
        tabs.insert(index, tab)


def add_field(paragraph, instruction: str, placeholder: str = "1"):
    """Insert a Word field (PAGE, NUMPAGES) that Word recalculates on render."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " " + instruction + " "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._element.append(node)
    return run


_VML_NS = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:v="urn:schemas-microsoft-com:vml" '
    'xmlns:o="urn:schemas-microsoft-com:office:office"'
)

_WATERMARK_XML = (
    "<w:p " + _VML_NS + "><w:r><w:pict>"
    '<v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" adj="10800"'
    ' path="m@7,l@8,m@5,21600l@11,21600e">'
    "<v:formulas>"
    '<v:f eqn="sum #0 0 10800"/><v:f eqn="prod #0 2 1"/><v:f eqn="sum 21600 0 @1"/>'
    '<v:f eqn="sum 0 0 @2"/><v:f eqn="sum 21600 0 @3"/><v:f eqn="if @0 @3 0"/>'
    '<v:f eqn="if @0 21600 @1"/><v:f eqn="if @0 0 @2"/><v:f eqn="if @0 @4 21600"/>'
    '<v:f eqn="mid @5 @6"/><v:f eqn="mid @8 @5"/><v:f eqn="mid @7 @8"/>'
    '<v:f eqn="mid @6 @7"/><v:f eqn="sum @6 0 @5"/>'
    "</v:formulas>"
    '<v:path textpathok="t" o:connecttype="custom"'
    ' o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800" o:connectangles="270,180,90,0"/>'
    '<v:textpath on="t" fitshape="t"/>'
    "</v:shapetype>"
    '<v:shape id="TechPulseWatermark" o:spid="_x0000_s2049" type="#_x0000_t136"'
    ' style="position:absolute;margin-left:0;margin-top:0;width:400pt;height:84pt;'
    "rotation:315;z-index:-251654144;mso-position-horizontal:center;"
    "mso-position-horizontal-relative:margin;mso-position-vertical:center;"
    'mso-position-vertical-relative:margin"'
    ' o:allowincell="f" fillcolor="#' + B.HEX_WATERMARK + '" stroked="f">'
    '<v:textpath style="font-family:&quot;' + B.FONT_BODY + '&quot;;font-weight:bold;'
    'v-text-align:center" string="' + B.WATERMARK_TEXT + '"/>'
    "</v:shape></w:pict></w:r></w:p>"
)


def add_watermark(header) -> None:
    """
    Diagonal watermark drawn from the page header so it sits behind body text on
    every page. Deliberately very light (see HEX_WATERMARK) so it never competes
    with what the student is reading.
    """
    anchor = header.add_paragraph()
    anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    anchor.paragraph_format.space_after = Pt(0)
    anchor._p.addnext(parse_xml(_WATERMARK_XML))


# ==========================================================================
# Document
# ==========================================================================

class TechPulseDocument:
    """Builder for a single branded learning document."""

    def __init__(
        self,
        *,
        title: str,
        subtitle: str,
        week_number: int,
        week_title: str,
        document_kind: str = "LEARNING GUIDE",
        summary_line: str = "",
    ) -> None:
        self.title = title
        self.subtitle = subtitle
        self.week_number = week_number
        self.week_title = week_title
        self.document_kind = document_kind
        self.summary_line = summary_line
        self.doc = Document()
        self.chapter_number = 0

        self._setup_page()
        self._setup_styles()
        self._build_cover()
        self._setup_header_footer()

    # ---------------------------------------------------------------- setup
    def _setup_page(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)      # A4
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.0)
        section.header_distance = Cm(1.1)
        section.footer_distance = Cm(1.0)
        section.different_first_page_header_footer = True

    def _apply_style_font(self, style) -> None:
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), B.FONT_BODY)

    def _setup_styles(self) -> None:
        styles = self.doc.styles

        normal = styles["Normal"]
        normal.font.size = B.SIZE_BODY
        normal.font.color.rgb = B.INK
        normal.font.name = B.FONT_BODY
        normal.paragraph_format.space_after = Pt(7)
        normal.paragraph_format.line_spacing = 1.18
        self._apply_style_font(normal)

        def tune(name, size, color, *, before, after):
            style = styles[name]
            style.font.size = size
            style.font.bold = True
            style.font.color.rgb = color
            style.font.name = B.FONT_BODY
            pf = style.paragraph_format
            pf.space_before = Pt(before)
            pf.space_after = Pt(after)
            pf.keep_with_next = True
            self._apply_style_font(style)

        tune("Heading 1", B.SIZE_H1, B.PRIMARY, before=20, after=8)
        tune("Heading 2", B.SIZE_H2, B.PRIMARY_GLOW, before=14, after=5)
        tune("Heading 3", B.SIZE_H3, B.INK, before=10, after=4)

    # ---------------------------------------------------------------- cover
    def _centred(self, text, size, color, *, bold=False, space_after=4):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(space_after)
        run = para.add_run(text)
        run.font.size = size
        run.font.bold = bold
        run.font.color.rgb = color
        set_run_font(run, B.FONT_BODY)
        return para

    def _build_cover(self) -> None:
        doc = self.doc
        if doc.paragraphs:
            doc.paragraphs[0].paragraph_format.space_after = Pt(0)
        doc.add_paragraph()

        self._centred(B.ORGANISATION, Pt(13), B.PRIMARY, bold=True, space_after=2)
        self._centred(B.PUBLICATION, Pt(11), B.ACCENT, bold=True, space_after=3)
        rule = self._centred(B.PROGRAM_NAME, Pt(10), B.MUTED, space_after=20)
        paragraph_border(rule, edges=("bottom",), hex_color=B.HEX_RULE, size=8, space=8)

        self._centred(self.document_kind, Pt(11), B.MUTED, bold=True, space_after=10)
        self._centred(self.title, B.SIZE_COVER_TITLE, B.PRIMARY, bold=True, space_after=6)
        self._centred(self.subtitle, Pt(13), B.INK, space_after=22)

        badge = self._centred(
            "WEEK {:02d}   -   {}".format(self.week_number, self.week_title.upper()),
            Pt(11),
            B.PRIMARY,
            bold=True,
            space_after=20,
        )
        badge.paragraph_format.space_before = Pt(6)
        shade_paragraph(badge, B.HEX_NOTE_BG)

        if self.summary_line:
            self._centred(self.summary_line, Pt(10.5), B.MUTED, space_after=26)

        for _ in range(9):
            doc.add_paragraph()

        self._centred("Prepared for:  Get Techy With Lucky  /  Tech Pulse Insider", Pt(10), B.MUTED, space_after=2)
        self._centred("Instructor:  " + B.INSTRUCTOR, Pt(10), B.MUTED, space_after=6)
        self._centred(B.TAGLINE, Pt(10), B.ACCENT, bold=True, space_after=0)

        doc.add_page_break()

    # -------------------------------------------------------- header/footer
    def _setup_header_footer(self) -> None:
        section = self.doc.sections[0]

        for para in section.first_page_header.paragraphs:
            para.text = ""
        for para in section.first_page_footer.paragraphs:
            para.text = ""

        header = section.header
        line1 = header.paragraphs[0]
        line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line1.paragraph_format.space_after = Pt(0)
        run = line1.add_run(B.ORGANISATION + "   |   " + B.PUBLICATION)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = B.PRIMARY
        set_run_font(run, B.FONT_BODY)

        line2 = header.add_paragraph()
        line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line2.paragraph_format.space_after = Pt(3)
        run = line2.add_run(
            "{} - WEEK {:02d}: {}".format(B.PROGRAM_NAME, self.week_number, self.week_title.upper())
        )
        run.font.size = B.SIZE_HEADER
        run.font.color.rgb = B.MUTED
        set_run_font(run, B.FONT_BODY)
        paragraph_border(line2, edges=("bottom",), hex_color=B.HEX_RULE, size=6, space=4)

        add_watermark(header)

        footer = section.footer
        para = footer.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        paragraph_border(para, edges=("top",), hex_color=B.HEX_RULE, size=6, space=6)

        usable = section.page_width - section.left_margin - section.right_margin
        para.paragraph_format.space_after = Pt(0)
        # The built-in Footer style ships Letter-width tab stops; drop them so the
        # page number lands on the A4 right margin instead of Word's default centre.
        para.paragraph_format.tab_stops.add_tab_stop(int(usable), WD_TAB_ALIGNMENT.RIGHT)
        clear_inherited_tabs(para)

        def foot_run(target, text, *, bold=False, color=B.MUTED, size=None):
            run = target.add_run(text)
            run.font.size = size or B.SIZE_FOOTER
            run.font.bold = bold
            run.font.color.rgb = color
            set_run_font(run, B.FONT_BODY)
            return run

        foot_run(para, B.FOOTER_IDENTITY, bold=True, color=B.PRIMARY)
        foot_run(para, "\t")
        foot_run(para, "Page ")
        page_field = add_field(para, "PAGE")
        foot_run(para, " of ")
        total_field = add_field(para, "NUMPAGES")
        for run in (page_field, total_field):
            run.font.size = B.SIZE_FOOTER
            run.font.color.rgb = B.MUTED
            set_run_font(run, B.FONT_BODY)

        tagline = footer.add_paragraph()
        tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tagline.paragraph_format.space_before = Pt(0)
        tagline.paragraph_format.space_after = Pt(0)
        foot_run(tagline, B.TAGLINE, bold=True, color=B.ACCENT, size=Pt(8))

    # ============================================================== content
    def para(self, text: str, *, italic=False, bold=False, size=None, color=None, space_after=7):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.size = size or B.SIZE_BODY
        run.font.italic = italic
        run.font.bold = bold
        run.font.color.rgb = color or B.INK
        set_run_font(run, B.FONT_BODY)
        return p

    def rich(self, segments, *, space_after=7):
        """segments: list of (text, {"bold": bool, "italic": bool, "code": bool})."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        for text, opts in segments:
            run = p.add_run(text)
            run.font.size = B.SIZE_CODE if opts.get("code") else B.SIZE_BODY
            run.font.bold = opts.get("bold", False)
            run.font.italic = opts.get("italic", False)
            run.font.color.rgb = opts.get("color", B.INK)
            set_run_font(run, B.FONT_CODE if opts.get("code") else B.FONT_BODY)
        return p

    def caption(self, text: str):
        return self.para(text, italic=True, size=B.SIZE_CAPTION, color=B.MUTED, space_after=9)

    def chapter(self, title: str):
        self.chapter_number += 1
        heading = self.doc.add_heading(
            "Chapter {} - {}".format(self.chapter_number, title), level=1
        )
        paragraph_border(heading, edges=("bottom",), hex_color=B.HEX_RULE, size=8, space=4)
        return heading

    def unnumbered_chapter(self, title: str):
        heading = self.doc.add_heading(title, level=1)
        paragraph_border(heading, edges=("bottom",), hex_color=B.HEX_RULE, size=8, space=4)
        return heading

    def section_heading(self, title: str):
        return self.doc.add_heading(title, level=2)

    def subsection(self, title: str):
        return self.doc.add_heading(title, level=3)

    def bullets(self, items, *, style="List Bullet"):
        for item in items:
            p = self.doc.add_paragraph(style=style)
            p.paragraph_format.space_after = Pt(3)
            if isinstance(item, tuple):
                lead, rest = item
                run = p.add_run(lead)
                run.font.bold = True
                run.font.size = B.SIZE_BODY
                run.font.color.rgb = B.INK
                set_run_font(run, B.FONT_BODY)
                run = p.add_run(rest)
            else:
                run = p.add_run(item)
            run.font.size = B.SIZE_BODY
            run.font.color.rgb = B.INK
            set_run_font(run, B.FONT_BODY)

    def numbered(self, items):
        self.bullets(items, style="List Number")

    def code_block(self, code: str, *, caption: str = ""):
        lines = code.strip("\n").split("\n")
        for index, line in enumerate(lines):
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_after = Pt(0)
            pf.space_before = Pt(6) if index == 0 else Pt(0)
            pf.line_spacing = 1.0
            pf.left_indent = Cm(0.4)
            run = p.add_run(line if line else " ")
            run.font.size = B.SIZE_CODE
            run.font.color.rgb = B.INK
            set_run_font(run, B.FONT_CODE)
            shade_paragraph(p, B.HEX_CODE_BG)
            edges = ["left", "right"]
            if index == 0:
                edges.append("top")
            if index == len(lines) - 1:
                edges.append("bottom")
            paragraph_border(p, edges=tuple(edges), hex_color=B.HEX_CODE_BORDER, size=4, space=2)
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)
        spacer.paragraph_format.space_before = Pt(0)
        if caption:
            self.caption(caption)

    _CALLOUTS = {
        "note": ("NOTE", B.HEX_NOTE_BG, B.PRIMARY),
        "why": ("WHY THIS MATTERS", B.HEX_WHY_BG, B.PRIMARY),
        "warning": ("COMMON MISTAKE", B.HEX_WARN_BG, B.ACCENT),
        "tip": ("BEST PRACTICE", B.HEX_TIP_BG, B.PRIMARY),
    }

    def callout(self, kind: str, text: str, *, title: str = ""):
        label, fill, color = self._CALLOUTS[kind]
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(8)
        pf.space_after = Pt(2)
        pf.left_indent = Cm(0.2)
        run = p.add_run(title.upper() if title else label)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = color
        set_run_font(run, B.FONT_BODY)
        shade_paragraph(p, fill)
        paragraph_border(p, edges=("left", "top", "right"), hex_color=B.HEX_RULE, size=4, space=3)

        body = self.doc.add_paragraph()
        bpf = body.paragraph_format
        bpf.space_before = Pt(0)
        bpf.space_after = Pt(9)
        bpf.left_indent = Cm(0.2)
        run = body.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = B.INK
        set_run_font(run, B.FONT_BODY)
        shade_paragraph(body, fill)
        paragraph_border(body, edges=("left", "bottom", "right"), hex_color=B.HEX_RULE, size=4, space=3)

    # --------------------------------------------------------------- tables
    def table(self, headers, rows, *, widths=None, font_size=None):
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        size = font_size or B.SIZE_TABLE

        for index, text in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            run = para.add_run(text)
            run.font.bold = True
            run.font.size = size
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_run_font(run, B.FONT_BODY)
            shade_cell(cell, B.HEX_TABLE_HEAD)

        # Repeat the header row when a long table spans a page break.
        header_props = table.rows[0]._tr.get_or_add_trPr()
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        header_props.append(repeat)

        for row_index, row in enumerate(rows):
            cells = table.add_row().cells
            for col_index, value in enumerate(row):
                cell = cells[col_index]
                cell.text = ""
                para = cell.paragraphs[0]
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.space_before = Pt(2)
                run = para.add_run(str(value))
                run.font.size = size
                run.font.color.rgb = B.INK
                set_run_font(run, B.FONT_BODY)
                if row_index % 2 == 1:
                    shade_cell(cell, B.HEX_TABLE_ZEBRA)

        if widths:
            for row in table.rows:
                for index, width in enumerate(widths):
                    row.cells[index].width = Cm(width)

        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        return table

    # -------------------------------------------------------- front matter
    def course_information(self, rows):
        self.unnumbered_chapter("Course Information")
        self.table(["Field", "Detail"], rows, widths=[4.6, 11.8])

    def learning_objectives(self, objectives):
        self.unnumbered_chapter("Learning Objectives")
        self.para("By the end of this week you should be able to:")
        self.numbered(objectives)

    def prerequisites(self, items):
        self.unnumbered_chapter("Prerequisites")
        self.para("Before starting this week, make sure you can already do the following:")
        self.bullets(items)

    def table_of_contents(self, entries):
        self.unnumbered_chapter("Table of Contents")
        self.caption(
            "This guide is written to be read in order. Each chapter builds on the one before it."
        )
        for index, entry in enumerate(entries, start=1):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.3)
            run = p.add_run("{}.  ".format(index))
            run.font.bold = True
            run.font.color.rgb = B.PRIMARY
            run.font.size = B.SIZE_BODY
            set_run_font(run, B.FONT_BODY)
            run = p.add_run(entry)
            run.font.size = B.SIZE_BODY
            run.font.color.rgb = B.INK
            set_run_font(run, B.FONT_BODY)
        self.page_break()

    def page_break(self):
        self.doc.add_page_break()

    # ---------------------------------------------------------- terminology
    def terminology(self, terms, *, intro=""):
        """terms: list of (term, definition, simple explanation, example)."""
        if intro:
            self.para(intro)
        self.table(
            ["Term", "Definition", "In Plain Words", "Example / Use Case"],
            terms,
            widths=[2.9, 5.2, 4.3, 4.0],
            font_size=Pt(9),
        )

    # ---------------------------------------------------------------- quiz
    def quiz(self, questions, *, instructions="", pass_mark="70%"):
        if instructions:
            self.para(instructions)
        self.rich([
            ("Pass mark: ", {"bold": True}),
            (pass_mark, {}),
            ("   |   Answers are at the end of this chapter.", {"italic": True, "color": B.MUTED}),
        ])
        for index, question in enumerate(questions, start=1):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(9)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run("Q{}. ".format(index))
            run.font.bold = True
            run.font.color.rgb = B.PRIMARY
            run.font.size = B.SIZE_BODY
            set_run_font(run, B.FONT_BODY)
            run = p.add_run(question["question"])
            run.font.size = B.SIZE_BODY
            run.font.color.rgb = B.INK
            set_run_font(run, B.FONT_BODY)

            kind = question.get("type", "mcq")
            if kind == "mcq":
                for key, option in question["options"]:
                    op = self.doc.add_paragraph()
                    op.paragraph_format.left_indent = Cm(0.8)
                    op.paragraph_format.space_after = Pt(1)
                    run = op.add_run("{})  {}".format(key, option))
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = B.INK
                    set_run_font(run, B.FONT_BODY)
            elif kind == "true_false":
                op = self.doc.add_paragraph()
                op.paragraph_format.left_indent = Cm(0.8)
                op.paragraph_format.space_after = Pt(1)
                run = op.add_run("True  /  False")
                run.font.size = Pt(10.5)
                run.font.color.rgb = B.INK
                set_run_font(run, B.FONT_BODY)
            else:
                op = self.doc.add_paragraph()
                op.paragraph_format.left_indent = Cm(0.8)
                op.paragraph_format.space_after = Pt(1)
                run = op.add_run("Short answer - write two or three sentences.")
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = B.MUTED
                set_run_font(run, B.FONT_BODY)

        self.section_heading("Answer Key")
        self.caption("Mark your own work honestly - the goal is to find your gaps, not to score well.")
        rows = []
        for index, question in enumerate(questions, start=1):
            rows.append(["Q{}".format(index), question["answer"], question.get("why", "")])
        self.table(["#", "Answer", "Why"], rows, widths=[1.2, 4.0, 11.2], font_size=Pt(9.5))

    # ---------------------------------------------------------- assignment
    def assignment(self, spec):
        self.section_heading("Objective")
        self.para(spec["objective"])

        self.section_heading("Scenario")
        self.para(spec["scenario"])

        self.section_heading("Requirements")
        self.bullets(spec["requirements"])

        self.section_heading("Expected Output")
        self.para(spec["expected_output"])

        self.section_heading("Technical Requirements")
        self.bullets(spec["technical_requirements"])

        self.section_heading("Submission Instructions")
        self.numbered(spec["submission"])

        self.section_heading("Evaluation Criteria")
        self.table(["Criterion", "What Earns Full Marks", "Weight"], spec["evaluation"],
                   widths=[4.2, 9.8, 2.4], font_size=Pt(9.5))

        if spec.get("bonus"):
            self.section_heading("Bonus Challenges")
            self.bullets(spec["bonus"])

    # -------------------------------------------------------------- outro
    def summary(self, points):
        self.bullets(points)

    def references(self, items):
        self.unnumbered_chapter("Further Reading and References")
        self.caption(
            "Documentation changes faster than any printed guide. Learning to read the official "
            "docs is part of the skill you are building."
        )
        for label, url in items:
            p = self.doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(label + " - ")
            run.font.bold = True
            run.font.size = B.SIZE_BODY
            run.font.color.rgb = B.INK
            set_run_font(run, B.FONT_BODY)
            run = p.add_run(url)
            run.font.size = Pt(10)
            run.font.color.rgb = B.PRIMARY_GLOW
            set_run_font(run, B.FONT_BODY)

    def save(self, path: str) -> str:
        self.doc.save(path)
        return path
